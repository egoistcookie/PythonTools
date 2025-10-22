import os
import tkinter as tk
from tkinter import filedialog, messagebox
import re
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# 尝试导入docx2pdf作为备用转换方法
try:
    import docx2pdf
    USE_DOCX2PDF = True
except ImportError:
    USE_DOCX2PDF = False
    print("docx2pdf库未安装，将使用reportlab作为备选转换方案")

class WordToPdfConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("Word转PDF转换器")
        self.root.geometry("500x300")
        self.root.resizable(False, False)
        
        # 设置中文字体
        self.font = ("SimHei", 10)
        
        # 初始化字体注册状态
        self.font_registered = False
        
        # 创建主框架
        self.main_frame = tk.Frame(root)
        self.main_frame.pack(pady=20, padx=20, fill=tk.BOTH, expand=True)
        
        # 创建标题
        self.title_label = tk.Label(self.main_frame, text="Word转PDF转换器", font=("SimHei", 16, "bold"))
        self.title_label.pack(pady=10)
        
        # 文件路径显示
        self.file_path_var = tk.StringVar()
        self.file_path_var.set("未选择文件")
        self.file_path_entry = tk.Entry(self.main_frame, textvariable=self.file_path_var, width=50, font=self.font)
        self.file_path_entry.pack(pady=10, fill=tk.X)
        
        # 选择文件按钮
        self.select_button = tk.Button(self.main_frame, text="选择Word文件", command=self.select_file, font=self.font)
        self.select_button.pack(pady=10)
        
        # 转换按钮
        self.convert_button = tk.Button(self.main_frame, text="转换为PDF", command=self.convert_to_pdf, font=self.font)
        self.convert_button.pack(pady=10)
        
        # 状态标签
        self.status_var = tk.StringVar()
        self.status_var.set("准备就绪")
        self.status_label = tk.Label(self.main_frame, textvariable=self.status_var, font=self.font)
        self.status_label.pack(pady=20)
    
    def select_file(self):
        file_path = filedialog.askopenfilename(
            title="选择Word文件",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if file_path:
            self.file_path_var.set(file_path)
            self.status_var.set(f"已选择文件: {os.path.basename(file_path)}")
    
    def convert_to_pdf(self):
        # 创建详细的调试日志文件
        debug_log_path = os.path.join(os.getcwd(), 'detailed_debug_log.txt')
        with open(debug_log_path, 'w', encoding='utf-8') as f:
            f.write("=== 详细转换调试日志 ===\n")
        
        def debug_log(message):
            """记录详细调试信息"""
            print(f"[调试] {message}")
            with open(debug_log_path, 'a', encoding='utf-8') as f:
                f.write(f"{message}\n")
        
        file_path = self.file_path_var.get()
        debug_log(f"开始转换文件: {file_path}")
        

        
        if file_path == "未选择文件":
            messagebox.showerror("错误", "请先选择一个Word文件")
            return
        
        if not os.path.exists(file_path):
            messagebox.showerror("错误", "所选文件不存在")
            return
        
        try:
            self.status_var.set("正在转换...")
            self.root.update()
            
            # 转换文件
            pdf_path = os.path.splitext(file_path)[0] + ".pdf"
            
            # 验证PDF路径是否可写
            pdf_dir = os.path.dirname(pdf_path)
            if not pdf_dir:
                pdf_dir = os.getcwd()
            
            if not os.access(pdf_dir, os.W_OK):
                messagebox.showerror("错误", f"没有写入权限: {pdf_dir}")
                self.status_var.set("准备就绪")
                return
            
            # 按照优先级顺序尝试转换方法
            conversion_success = False
            
            # 首先尝试使用docx2pdf
            if USE_DOCX2PDF:
                try:
                    debug_log("尝试使用docx2pdf转换")
                    self.status_var.set("使用docx2pdf转换...")
                    self.root.update()
                    
                    # 确保目标PDF路径存在
                    if os.path.exists(pdf_path):
                        debug_log(f"删除已存在的PDF文件: {pdf_path}")
                        os.remove(pdf_path)
                    
                    # 使用docx2pdf进行转换
                    docx2pdf.convert(file_path, pdf_path)
                    
                    # 验证PDF文件是否成功创建且不为空
                    if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                        debug_log(f"docx2pdf转换成功: {pdf_path}")
                        conversion_success = True
                    else:
                        debug_log("docx2pdf转换失败: PDF文件不存在或为空")
                except Exception as docx2pdf_error:
                    debug_log(f"docx2pdf转换失败: {str(docx2pdf_error)}")
                    # 如果有具体的错误信息，记录更多详情
                    if hasattr(docx2pdf_error, 'args'):
                        debug_log(f"docx2pdf错误详情: {docx2pdf_error.args}")
            else:
                debug_log("docx2pdf不可用(未安装)")
            
            # 如果docx2pdf失败，尝试使用WPS命令行
            if not conversion_success:
                try:
                    debug_log("尝试使用WPS命令行转换")
                    self.status_var.set("使用WPS转换...")
                    self.root.update()
                    
                    # 查找WPS安装位置
                    wps_path = None
                    possible_wps_paths = [
                        "C:\\Program Files\\Kingsoft\\WPS Office\\11.2.0.12841\\office6\\wps.exe",
                        "C:\\Program Files (x86)\\Kingsoft\\WPS Office\\11.2.0.12841\\office6\\wps.exe",
                        "C:\\Program Files\\Kingsoft\\WPS Office\\office6\\wps.exe",
                        "C:\\Program Files (x86)\\Kingsoft\\WPS Office\\office6\\wps.exe"
                    ]
                    
                    for path in possible_wps_paths:
                        if os.path.exists(path):
                            wps_path = path
                            break
                    
                    if wps_path:
                        # 构建WPS命令行
                        import subprocess
                        cmd = [wps_path, "-convert", pdf_path, file_path]
                        subprocess.run(cmd, check=True, timeout=60)
                        
                        # 验证PDF文件是否成功创建且不为空
                        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                            debug_log(f"WPS转换成功: {pdf_path}")
                            conversion_success = True
                        else:
                            debug_log("WPS转换失败: PDF文件不存在或为空")
                    else:
                        debug_log("WPS未找到")
                except Exception as wps_error:
                    debug_log(f"WPS转换失败: {str(wps_error)}")
            
            # 如果前两种方法都失败，使用reportlab
            if not conversion_success:
                debug_log("尝试使用reportlab转换")
                self.status_var.set("使用reportlab转换...")
                self.root.update()
            else:
                # 如果前面的转换成功，直接返回
                self.status_var.set(f"转换成功! PDF已保存至: {os.path.basename(pdf_path)}")
                messagebox.showinfo("成功", f"文件已成功转换为PDF\n保存路径: {pdf_path}")
                return
            
            # 仅在reportlab方法中进行后续处理
            if not conversion_success:
                debug_log("开始执行reportlab转换流程")
                # 注册中文字体以支持中文显示
                font_registered = self.register_chinese_fonts()
                base_font = 'SimSun' if font_registered else 'Helvetica'
                debug_log(f"字体注册状态: {font_registered}, 基础字体: {base_font}")
                
                # 使用文档原有颜色设置
                debug_log("使用文档原有颜色设置")
                
                # 使用python-docx读取Word文档
                try:
                    doc = Document(file_path)
                    debug_log(f"成功读取Word文档，包含{len(doc.paragraphs)}个段落和{len(doc.tables)}个表格")
                except Exception as doc_error:
                    debug_log(f"读取Word文档失败: {str(doc_error)}")
                    raise Exception(f"读取Word文档失败: {str(doc_error)}")
                
                # 创建PDF文档 - 添加文件对象参数以避免'write'属性错误
                try:
                    pdf = SimpleDocTemplate(pdf_path, pagesize=A4)
                    debug_log(f"成功创建PDF模板: {pdf_path}")
                except Exception as template_error:
                    error_msg = f"创建PDF模板失败: {str(template_error)}"
                    debug_log(error_msg)
                    raise Exception(error_msg)
            else:
                # 如果已经通过其他方法转换成功，直接跳过reportlab部分
                debug_log("已通过其他方法转换成功，跳过reportlab部分")
                return
            
            # 获取样式
            styles = getSampleStyleSheet()
            
            # 创建基本样式，设置中文字体并启用HTML
            custom_styles = {}
            
            # 导入用于处理HTML的模块
            from reportlab.lib.enums import TA_LEFT
            from reportlab.lib.colors import Color
            custom_styles['Normal'] = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName='SimSun' if font_registered else 'Helvetica',
                fontSize=12,
                leading=15,
                # 启用HTML处理以支持链接
                wordWrap='CJK',
                )
            
            # 创建各级标题样式，设置中文字体
            for i in range(1, 7):
                heading_style = ParagraphStyle(
                    f'CustomHeading{i}',
                    parent=styles.get(f'Heading{i}', styles['Heading1']),
                    fontName='SimHei' if font_registered else 'Helvetica-Bold',
                    fontSize=16 - i * 2 if i <= 4 else 10,
                    leading=20 - i * 2 if i <= 4 else 12
                )
                custom_styles[f'Heading{i}'] = heading_style
            
            # 创建目录样式
            custom_styles['TOC'] = styles['Normal']
            
            # 为目录项创建不同级别的样式，设置中文字体
            for i in range(1, 5):
                toc_style = ParagraphStyle(
                    f'CustomTOC{i}',
                    parent=custom_styles['Normal'],
                    leftIndent=i * 30,
                    spaceAfter=6,
                    fontName='SimSun' if font_registered else 'Helvetica'
                )
                custom_styles[f'TOC{i}'] = toc_style
            
            # 创建内容列表
            flowables = []
            
            # 收集文档中的标题，用于重建目录
            headings = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    level_match = re.search(r'\d+', para.style.name)
                    if level_match:
                        level = int(level_match.group())
                        headings.append((level, para.text))
            
            # 检查是否有目录内容
            has_toc = any('目录' in para.text or 'Contents' in para.text for para in doc.paragraphs)
            
            # 如果文档中提到了目录，添加重建的目录
            if has_toc:
                # 添加目录标题
                toc_title = Paragraph("目录", custom_styles.get('Heading1', custom_styles['Normal']))
                flowables.append(toc_title)
                flowables.append(Spacer(1, 0.3*inch))
                
                # 添加目录项
                for level, text in headings:
                    if level <= 4:  # 只包含前4级标题
                        toc_text = text
                        toc_item = Paragraph(toc_text, custom_styles.get(f'TOC{level}', custom_styles['TOC']))
                        flowables.append(toc_item)
                
                flowables.append(Spacer(1, 0.5*inch))
            
            # 处理每个段落并保留基本格式
            for para in doc.paragraphs:
                # 跳过目录占位文本（如果有的话）
                if has_toc and ('目录' in para.text or 'Contents' in para.text) and para.style.name.startswith('Heading'):
                    continue
                
                if para.text.strip():
                    # 简化样式处理，避免复杂的格式转换
                    if para.style.name in custom_styles:
                        pdf_style = custom_styles[para.style.name]
                    elif para.style.name.startswith('Heading'):
                        level_match = re.search(r'\d+', para.style.name)
                        if level_match:
                            level = int(level_match.group())
                            if level <= 6:
                                pdf_style = custom_styles.get(f'Heading{level}', custom_styles['Normal'])
                            else:
                                pdf_style = custom_styles['Normal']
                        else:
                            pdf_style = custom_styles['Normal']
                    else:
                        pdf_style = custom_styles['Normal']
                    
                    # 处理文本和颜色
                    formatted_text = self.process_text_with_formatting(para)
                    formatted_text = formatted_text.replace('\t', '    ')
                    
                    # 添加段落，捕获可能的编码问题
                    try:
                        # 创建段落时启用HTML解析以支持链接
                        flowables.append(Paragraph(formatted_text, pdf_style))
                        
                        # 添加适当的间距
                        if para.style.name.startswith('Heading'):
                            flowables.append(Spacer(1, 0.2*inch))
                    except Exception as para_error:
                        # 如果段落处理失败，尝试简化文本
                        try:
                            simplified_text = ''.join(char for char in formatted_text if ord(char) < 128 or char in '，。；：、？！""''（）【】《》')
                            flowables.append(Paragraph(simplified_text, pdf_style))
                        except:
                            # 如果仍然失败，添加错误标记
                            flowables.append(Paragraph("[无法转换的文本]", pdf_style))
            
            # 简化表格处理
            for table in doc.tables:
                try:
                    data = []
                    # 获取表格数据
                    for row in table.rows:
                        data_row = []
                        for cell in row.cells:
                            # 简化单元格文本处理
                            cell_text = cell.text.strip()
                            data_row.append(cell_text)
                        data.append(data_row)
                    
                    if data:
                        # 创建表格
                        table_obj = Table(data)
                        # 使用基本样式，不再强制设置文本颜色为红色
                        table_style = TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            # 移除强制的TEXTCOLOR设置，使用原有颜色
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('FONTNAME', (0, 0), (-1, -1), 'SimSun' if font_registered else 'Helvetica'),  # 使用中文字体
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                            ('TOPPADDING', (0, 0), (-1, -1), 8),
                            ('LEFTPADDING', (0, 0), (-1, -1), 10),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ])
                        table_obj.setStyle(table_style)
                        flowables.append(table_obj)
                        flowables.append(Spacer(1, 0.5*inch))
                except Exception as table_error:
                    # 表格处理失败时添加错误标记
                    error_para = Paragraph(f"[无法转换的表格: {str(table_error)[:50]}...]", custom_styles['Normal'])
                    flowables.append(error_para)
            
            # 构建PDF - 加强错误处理
            try:

                
                # 确保flowables不为空
                if not flowables:
                    flowables.append(Paragraph("[空文档]", custom_styles['Normal']))
                
                pdf.build(flowables)
                
                # 验证PDF文件是否成功创建且不为空
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    self.status_var.set(f"转换成功! PDF已保存至: {os.path.basename(pdf_path)}")
                    messagebox.showinfo("成功", f"文件已成功转换为PDF\n保存路径: {pdf_path}")
                else:
                    raise Exception(f"PDF文件创建失败或为空: {pdf_path}")
            except Exception as build_error:
                # 详细记录错误信息
                error_msg = f"PDF生成失败: {str(build_error)}"
                
                # 检查是否是文件写入问题
                if "write" in str(build_error).lower() or "NoneType" in str(build_error):
                    error_msg += "\n\n可能原因:\n1. 没有写入权限\n2. 文件被其他程序占用\n3. 磁盘空间不足\n4. 路径包含特殊字符\n5. PDF构建器初始化失败"
                
                raise Exception(error_msg)
            
        except Exception as e:
            error_info = str(e)
            debug_log(f"转换失败: {error_info}")
            self.status_var.set(f"转换失败: {error_info}")
            messagebox.showerror("错误", f"转换过程中发生错误:\n{error_info}")

    def process_text_with_formatting(self, para):
        """
        处理文本格式并转义特殊字符
        """
        result = []
        
        # 处理段落中的每个run
        if len(para.runs) > 0:
            for run in para.runs:
                # 转义HTML特殊字符
                text = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                result.append(text)
            return ''.join(result)
        
        # 如果没有runs，返回纯文本
        return para.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    
    def register_chinese_fonts(self):
        """
        尝试注册中文字体以支持中文显示
        返回True表示字体注册成功，False表示失败
        """
        if self.font_registered:
            return True
        
        try:
            # 尝试从系统字体目录加载字体
            font_paths = [
                "C:\\Windows\\Fonts\\simsun.ttc",  # 宋体
                "C:\\Windows\\Fonts\\simhei.ttf",  # 黑体
                "C:\\WINNT\\Fonts\\simsun.ttc",
                "C:\\WINNT\\Fonts\\simhei.ttf"
            ]
            
            # 注册宋体（用于正文）
            simsun_path = None
            for path in font_paths:
                if path.endswith("simsun.ttc") and os.path.exists(path):
                    simsun_path = path
                    break
            
            # 注册黑体（用于标题）
            simhei_path = None
            for path in font_paths:
                if path.endswith("simhei.ttf") and os.path.exists(path):
                    simhei_path = path
                    break
            
            # 尝试注册找到的字体
            if simsun_path:
                pdfmetrics.registerFont(TTFont('SimSun', simsun_path))
            if simhei_path:
                pdfmetrics.registerFont(TTFont('SimHei', simhei_path))
            
            # 如果至少注册了一个字体，则认为成功
            self.font_registered = simsun_path is not None or simhei_path is not None
            return self.font_registered
            
        except Exception as e:
            print(f"字体注册失败: {str(e)}")
            self.font_registered = False
            return False

if __name__ == "__main__":
    root = tk.Tk()
    app = WordToPdfConverter(root)
    root.mainloop()